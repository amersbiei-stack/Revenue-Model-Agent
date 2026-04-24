"""Convert spec.md to spec.docx using python-docx.

Handles: H1-H4 headers, paragraphs, GFM-style tables, fenced code blocks,
unordered/ordered lists, bold/italic inline, inline code.
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = Path(r"C:\Users\amers\Downloads\Revenue Model\spec.md")
DST = Path(r"C:\Users\amers\Downloads\Revenue Model\spec.docx")


def set_cell_shading(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def add_inline(paragraph, text, base_style=None):
    """Parse inline markdown: **bold**, *italic*, `code`, [text](url)."""
    # Tokenize — keep delimiters
    pattern = re.compile(
        r"(\*\*[^\*]+\*\*|__[^_]+__|\*[^\*]+\*|_[^_]+_|`[^`]+`|\[[^\]]+\]\([^\)]+\))"
    )
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            if base_style == "code":
                run.font.name = "Consolas"
                run.font.size = Pt(9)
        tok = m.group(0)
        if tok.startswith("**") and tok.endswith("**"):
            r = paragraph.add_run(tok[2:-2])
            r.bold = True
        elif tok.startswith("__") and tok.endswith("__"):
            r = paragraph.add_run(tok[2:-2])
            r.bold = True
        elif tok.startswith("*") and tok.endswith("*"):
            r = paragraph.add_run(tok[1:-1])
            r.italic = True
        elif tok.startswith("_") and tok.endswith("_"):
            r = paragraph.add_run(tok[1:-1])
            r.italic = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = paragraph.add_run(tok[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        elif tok.startswith("["):
            m2 = re.match(r"\[([^\]]+)\]\(([^\)]+)\)", tok)
            if m2:
                r = paragraph.add_run(m2.group(1))
                r.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
                r.underline = True
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        if base_style == "code":
            run.font.name = "Consolas"
            run.font.size = Pt(9)


def is_table_sep(line):
    s = line.strip()
    if not s.startswith("|"):
        return False
    # remove edge pipes
    inner = s.strip("|")
    cells = [c.strip() for c in inner.split("|")]
    return all(re.match(r"^:?-{3,}:?$", c) for c in cells) and len(cells) >= 2


def split_table_row(line):
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def add_table(doc, header_cells, rows):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(header_cells))
    tbl.style = "Light Grid Accent 1"
    tbl.autofit = True
    # Header
    for j, h in enumerate(header_cells):
        cell = tbl.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        add_inline(p, h)
        for run in p.runs:
            run.bold = True
        set_cell_shading(cell, "305496")
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Rows
    for i, row in enumerate(rows, start=1):
        for j in range(len(header_cells)):
            cell = tbl.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            val = row[j] if j < len(row) else ""
            add_inline(p, val)
    doc.add_paragraph()


def add_code_block(doc, lines, lang=""):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    # light gray background via shading on paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    text = "\n".join(lines)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def convert(md_text, doc):
    lines = md_text.split("\n")
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Fenced code block
        fence = re.match(r"^```(\w*)\s*$", stripped)
        if fence:
            lang = fence.group(1)
            buf = []
            i += 1
            while i < n and not re.match(r"^```\s*$", lines[i].strip()):
                buf.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            add_code_block(doc, buf, lang)
            continue

        # Horizontal rule
        if re.match(r"^-{3,}$|^\*{3,}$|^_{3,}$", stripped):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bot = OxmlElement("w:bottom")
            bot.set(qn("w:val"), "single")
            bot.set(qn("w:sz"), "6")
            bot.set(qn("w:space"), "1")
            bot.set(qn("w:color"), "999999")
            pBdr.append(bot)
            pPr.append(pBdr)
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            h = doc.add_heading(level=min(level, 4))
            add_inline(h, text)
            i += 1
            continue

        # Table: a line starting with | followed by a separator line
        if stripped.startswith("|") and i + 1 < n and is_table_sep(lines[i + 1]):
            header = split_table_row(line)
            i += 2
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append(split_table_row(lines[i]))
                i += 1
            add_table(doc, header, rows)
            continue

        # Blockquote
        if stripped.startswith(">"):
            text = stripped.lstrip(">").strip()
            p = doc.add_paragraph(style="Intense Quote")
            add_inline(p, text)
            i += 1
            continue

        # Bullet list
        bm = re.match(r"^(\s*)[-*+]\s+(.+?)\s*$", line)
        if bm:
            indent_spaces = len(bm.group(1))
            level = min(indent_spaces // 2, 4)
            text = bm.group(2)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
            add_inline(p, text)
            i += 1
            continue

        # Numbered list
        nm = re.match(r"^(\s*)(\d+)\.\s+(.+?)\s*$", line)
        if nm:
            indent_spaces = len(nm.group(1))
            text = nm.group(3)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Inches(0.25 + 0.25 * (indent_spaces // 2))
            add_inline(p, text)
            i += 1
            continue

        # Blank
        if stripped == "":
            i += 1
            continue

        # Regular paragraph — collect consecutive non-special lines
        para_lines = [line]
        i += 1
        while i < n:
            nxt = lines[i]
            nstr = nxt.strip()
            if nstr == "":
                break
            if re.match(r"^#{1,6}\s+", nxt):
                break
            if re.match(r"^```", nstr):
                break
            if nstr.startswith("|"):
                break
            if re.match(r"^(\s*)[-*+]\s+", nxt):
                break
            if re.match(r"^(\s*)\d+\.\s+", nxt):
                break
            if nstr.startswith(">"):
                break
            if re.match(r"^-{3,}$|^\*{3,}$|^_{3,}$", nstr):
                break
            para_lines.append(nxt)
            i += 1
        text = " ".join(l.strip() for l in para_lines)
        p = doc.add_paragraph()
        add_inline(p, text)


def main():
    md = SRC.read_text(encoding="utf-8")
    doc = Document()

    # Base style tweaks
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Tight margins for readability
    for section in doc.sections:
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)

    convert(md, doc)
    doc.save(DST)
    print(f"Wrote {DST}")
    print(f"Size: {DST.stat().st_size:,} bytes")


if __name__ == "__main__":
    main()
